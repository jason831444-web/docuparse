from __future__ import annotations

import csv
import json
import textwrap
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SPEC = ROOT / "eval" / "specs" / "eval_documents.json"
DEFAULT_CORPUS = ROOT / "eval" / "corpus"


def load_spec(spec_path: Path = DEFAULT_SPEC) -> dict:
    return json.loads(spec_path.read_text(encoding="utf-8"))


def generate_corpus(spec_path: Path = DEFAULT_SPEC, corpus_dir: Path = DEFAULT_CORPUS) -> list[Path]:
    spec = load_spec(spec_path)
    corpus_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for item in spec["documents"]:
        path = corpus_dir / item["filename"]
        _write_document(path, item["format"], item["content"])
        written.append(path)
    return written


def _write_document(path: Path, doc_format: str, content: dict) -> None:
    if doc_format == "pdf":
        _write_pdf(path, content)
        return
    if doc_format == "docx":
        _write_docx(path, content)
        return
    if doc_format in {"txt", "md"}:
        _write_text(path, content)
        return
    if doc_format == "png":
        _write_image(path, content)
        return
    if doc_format == "json":
        path.write_text(json.dumps(content["json_object"], indent=2), encoding="utf-8")
        return
    if doc_format == "csv":
        _write_csv(path, content)
        return
    if doc_format == "xlsx":
        _write_xlsx(path, content)
        return
    if doc_format == "html":
        _write_html(path, content)
        return
    if doc_format == "xml":
        _write_xml(path, content)
        return
    raise ValueError(f"Unsupported eval document format: {doc_format}")


def _write_pdf(path: Path, content: dict) -> None:
    doc = fitz.open()
    title = content.get("title")
    lines = list(content.get("lines", []))
    page = doc.new_page(width=612, height=792)
    y = 52
    if title:
        page.insert_textbox(fitz.Rect(48, y, 560, y + 42), title, fontsize=18, fontname="helv")
        y += 44
    for line in lines:
        chunks = textwrap.wrap(line, width=92) or [""]
        block_height = max(22, 15 * len(chunks) + 10)
        if y + block_height > 740:
            page = doc.new_page(width=612, height=792)
            y = 52
        page.insert_textbox(
            fitz.Rect(48, y, 560, y + block_height),
            "\n".join(chunks),
            fontsize=11.5,
            fontname="helv",
            lineheight=1.25,
        )
        y += block_height
    doc.save(path)
    doc.close()


def _write_docx(path: Path, content: dict) -> None:
    doc = DocxDocument()
    title = content.get("title")
    if title:
        doc.add_heading(title, level=1)
    for line in content.get("lines", []):
        if not line.strip():
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(path)


def _write_text(path: Path, content: dict) -> None:
    pieces = []
    if content.get("title"):
        pieces.append(content["title"])
    pieces.extend(content.get("lines", []))
    path.write_text("\n".join(pieces).strip() + "\n", encoding="utf-8")


def _write_image(path: Path, content: dict) -> None:
    image = Image.new("RGB", (1200, 1400), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("Arial.ttf", 34)
        body_font = ImageFont.truetype("Arial.ttf", 28)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()

    y = 60
    title = content.get("title")
    if title:
        draw.text((60, y), title, fill="black", font=title_font)
        y += 70
    for line in content.get("lines", []):
        wrapped = textwrap.wrap(line, width=40) or [""]
        for chunk in wrapped:
            draw.text((60, y), chunk, fill="black", font=body_font)
            y += 42
        y += 8
    if content.get("ocr_noise"):
        for x in range(90, 1120, 170):
            draw.line((x, 120, x + 70, 1280), fill=(225, 225, 225), width=2)
        for y_line in range(230, 1250, 190):
            draw.line((45, y_line, 1150, y_line + 20), fill=(230, 230, 230), width=2)
    image.save(path)


def _write_csv(path: Path, content: dict) -> None:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        for row in content.get("rows", []):
            writer.writerow(row)


def _write_xlsx(path: Path, content: dict) -> None:
    workbook = Workbook()
    first = True
    for sheet_spec in content.get("sheets", []):
        if first:
            sheet = workbook.active
            sheet.title = sheet_spec["name"]
            first = False
        else:
            sheet = workbook.create_sheet(title=sheet_spec["name"])
        for row in sheet_spec.get("rows", []):
            sheet.append(row)
    workbook.save(path)


def _write_html(path: Path, content: dict) -> None:
    title = content.get("title", "Evaluation Document")
    paragraphs = "\n".join(f"<p>{line}</p>" for line in content.get("lines", []))
    html = f"""<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <title>{title}</title>
  </head>
  <body>
    <h1>{title}</h1>
    {paragraphs}
  </body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def _write_xml(path: Path, content: dict) -> None:
    title = content.get("title", "Evaluation Document")
    body = "\n".join(f"  <line>{line}</line>" for line in content.get("lines", []))
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<document>
  <title>{title}</title>
{body}
</document>
"""
    path.write_text(xml, encoding="utf-8")


if __name__ == "__main__":
    paths = generate_corpus()
    print(f"Generated {len(paths)} evaluation documents in {DEFAULT_CORPUS}")
